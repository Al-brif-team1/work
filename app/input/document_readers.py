"""Извлечение текста из офисных форматов документов для входного модуля брифа."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError


def read_docx(path: Path) -> str:
    """Извлекает текст из .docx: параграфы и содержимое таблиц по порядку."""
    try:
        document = Document(str(path))
    except PackageNotFoundError as exc:
        raise ValueError(f"File is not a valid .docx package: {path}") from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    text = "\n".join(parts)
    if not text.strip():
        raise ValueError(f"No extractable text found in .docx file: {path}")

    return text